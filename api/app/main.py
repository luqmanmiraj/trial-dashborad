import os
import base64
import mimetypes
import subprocess
from typing import List
from datetime import datetime, date, timedelta

from fastapi import FastAPI, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel

from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
try:
    import boto3
    from botocore.exceptions import BotoCoreError, ClientError
except Exception:  # boto3 optional for local use
    boto3 = None
    BotoCoreError = ClientError = Exception

# ---------------- Config ----------------
PEN_EMAIL = os.getenv('PEN_EMAIL', 'office@pen.com')
EMAIL_ADDRESS = os.getenv('EMAIL_ADDRESS', 'your_email@example.com')
TOKEN_FILE = os.getenv('TOKEN_FILE', 'token.json')
SCOPES = ['https://www.googleapis.com/auth/gmail.send']

BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FINISHED_DIR = os.path.join(BASE_DIR, 'finished_noms')
os.makedirs(FINISHED_DIR, exist_ok=True)

MGO_TEMPLATE = os.getenv('MGO_TEMPLATE', os.path.join(BASE_DIR, 'mgo_nom_template.docx'))
IFO_TEMPLATE = os.getenv('IFO_TEMPLATE', os.path.join(BASE_DIR, 'ifo_nom_template.docx'))
BOTH_TEMPLATE = os.getenv('BOTH_TEMPLATE', os.path.join(BASE_DIR, 'mgo_ifo_nom_template.docx'))

# Invoice templates (will use nomination templates as fallback)
MGO_INVOICE_TEMPLATE = os.getenv('MGO_INVOICE_TEMPLATE', os.path.join(BASE_DIR, 'mgo_invoice_template.docx'))
IFO_INVOICE_TEMPLATE = os.getenv('IFO_INVOICE_TEMPLATE', os.path.join(BASE_DIR, 'ifo_invoice_template.docx'))
BOTH_INVOICE_TEMPLATE = os.getenv('BOTH_INVOICE_TEMPLATE', os.path.join(BASE_DIR, 'mgo_ifo_invoice_template.docx'))

LIBREOFFICE_PATH = os.getenv('LIBREOFFICE_PATH', r'C:\\Program Files\\LibreOffice\\program\\soffice.exe')


DISABLE_EMAIL = os.getenv('DISABLE_EMAIL', '0') == '1'

# S3 config (optional). If S3_BUCKET is set, generated files will be uploaded.
S3_BUCKET = os.getenv('S3_BUCKET')
S3_PREFIX = os.getenv('S3_PREFIX', 'noms/')
S3_REGION = os.getenv('AWS_REGION') or os.getenv('AWS_DEFAULT_REGION') or 'us-east-1'

_s3_client = None
def get_s3_client():
    global _s3_client
    if not S3_BUCKET:
        return None
    if boto3 is None:
        return None
    if _s3_client is None:
        _s3_client = boto3.client('s3', region_name=S3_REGION)
    return _s3_client

def authenticate():
    if DISABLE_EMAIL:
        return None
    if os.path.exists(TOKEN_FILE):
        creds = Credentials.from_authorized_user_file(TOKEN_FILE, SCOPES)
        return build('gmail', 'v1', credentials=creds)
    return None


def send_email(recipients, subject, body, attachments=None):
    service = authenticate()
    if service is None:
        # Local test mode: skip sending
        print('[LOCAL TEST] Email disabled or token missing. Skipping send.')
        return
    message = MIMEMultipart()
    message['To'] = ', '.join(recipients)
    message['Subject'] = subject
    message.attach(MIMEText(body, 'plain'))

    for file_path in attachments or []:
        if os.path.exists(file_path):
            filename = os.path.basename(file_path)
            mime_type, _ = mimetypes.guess_type(file_path)
            if mime_type is None:
                mime_type = 'application/octet-stream'
            main_type, sub_type = mime_type.split('/', 1)
            with open(file_path, 'rb') as f:
                part = MIMEBase(main_type, sub_type)
                part.set_payload(f.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename={filename}')
            message.attach(part)

    encoded_message = base64.urlsafe_b64encode(message.as_bytes()).decode('utf-8')
    service.users().messages().send(userId='me', body={'raw': encoded_message}).execute()


def replace_and_format_run(run, old_string, new_string):
    if old_string in run.text:
        run.text = run.text.replace(str(old_string), str(new_string))
        run.font.name = 'Nunito'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Tahoma')
        run.font.size = Pt(9)


def replace_strings_in_docx(doc_path, output_path, replacements, save):
    doc = Document(doc_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for old_string, new_string in replacements.items():
                replace_and_format_run(run, old_string, new_string)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for old_string, new_string in replacements.items():
                            replace_and_format_run(run, old_string, new_string)
    if save == 1:
        doc.save(output_path)


def get_bunker_date(value):
    if '-' in value:
        value = value.split('-')[1]
    value = value.split('.')
    return date(int(value[2]), int(value[1]), int(value[0]))


def convert_docx_to_pdf(input_path, output_path, librepath):
    if not input_path.endswith('.docx'):
        raise ValueError('Input file must be a .docx file.')
    if not output_path.endswith('.pdf'):
        raise ValueError('Output file must have a .pdf extension.')
    if not os.path.exists(input_path):
        raise FileNotFoundError(f'Template not found: {input_path}')

    output_dir = os.path.dirname(output_path)
    os.makedirs(output_dir, exist_ok=True)

    if not os.path.exists(librepath):
        # Fallback: skip conversion in local test
        print('[LOCAL TEST] LibreOffice not found. Returning DOCX instead of PDF.')
        return input_path
    command = [
        librepath,
        '--headless',
        '--convert-to', 'pdf',
        '--outdir', output_dir,
        input_path,
    ]
    try:
        subprocess.run(
            command,
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            timeout=25,  # fail fast if LibreOffice hangs
        )
        return output_path
    except Exception as e:
        # On any error or timeout, fall back to returning the original DOCX path
        print(f"[LOCAL TEST] Conversion skipped ({type(e).__name__}). Returning DOCX path.")
        return input_path


def delete_file(file_path):
    if os.path.exists(file_path):
        os.remove(file_path)


def upload_files_to_s3(file_paths):
    client = get_s3_client()
    if client is None:
        return []
    uploaded_urls = []
    for file_path in file_paths:
        try:
            filename = os.path.basename(file_path)
            key = f"{S3_PREFIX}{filename}"
            client.upload_file(file_path, S3_BUCKET, key)
            # Generate a presigned URL valid for 7 days
            url = client.generate_presigned_url(
                'get_object',
                Params={'Bucket': S3_BUCKET, 'Key': key},
                ExpiresIn=7 * 24 * 3600,
            )
            uploaded_urls.append({'key': key, 'url': url})
        except (BotoCoreError, ClientError) as e:
            print(f"[S3] Upload failed for {file_path}: {e}")
    return uploaded_urls


def process_both(vessel_name, vessel_imo, supply_dates, mgo_tons, mgo_price, ifo_tons, ifo_price, agent):
    global queued_up_files
    # Ensure template exists; if not, build a very simple one for local test
    if not os.path.exists(BOTH_TEMPLATE):
        tmp = Document()
        tmp.add_heading('Bunkering nomination (MGO + IFO)', level=1)
        tmp.add_paragraph(f'Vessel: {vessel_name} (IMO {vessel_imo})')
        tmp.add_paragraph(f'Date: {supply_dates}')
        tmp.save(BOTH_TEMPLATE)
    replacements = {
        'X1_CN': str('Simple Fuel FZCO').upper(),
        'X1_VSLN': str(vessel_name).upper(),
        'X1_IMO': vessel_imo,
        'X1_VSLSD': supply_dates,
        'X1_MQ': mgo_tons,
        'X1_MP': mgo_price - 2,
        'X1_IQ': ifo_tons,
        'X1_IP': ifo_price - 2,
        'X1_AGNT': agent,
    }
    replacements2 = {
        'X1_RN': str(get_bunker_date(replacements.get('X1_VSLSD')).strftime('%Y%m%d')) + '-NOM-' + str(replacements.get('X1_VSLN')).replace(' ', '_'),
        'X1_UC': 'USD',
        'X1_DDD': 30,
        'X1_DATE': get_bunker_date(replacements.get('X1_VSLSD')) - timedelta(days=10),
    }
    in_path = BOTH_TEMPLATE
    out_path = os.path.join(FINISHED_DIR, f"{replacements2['X1_RN']}.docx")
    replace_strings_in_docx(in_path, out_path, replacements, 1)
    replace_strings_in_docx(out_path, out_path, replacements2, 1)
    input_docx = out_path
    output_pdf = os.path.join(FINISHED_DIR, f"{replacements2['X1_RN']}.pdf")
    final_path = convert_docx_to_pdf(input_docx, output_pdf, LIBREOFFICE_PATH)
    if final_path.endswith('.pdf'):
        delete_file(out_path)
    queued_up_files.append(final_path)


def process_mgo(vessel_name, vessel_imo, supply_dates, mgo_tons, mgo_price, agent):
    global queued_up_files
    if not os.path.exists(MGO_TEMPLATE):
        tmp = Document()
        tmp.add_heading('Bunkering nomination (MGO)', level=1)
        tmp.add_paragraph(f'Vessel: {vessel_name} (IMO {vessel_imo})')
        tmp.add_paragraph(f'Date: {supply_dates}')
        tmp.save(MGO_TEMPLATE)
    replacements = {
        'X1_CN': str('Simple Fuel FZCO').upper(),
        'X1_VSLN': str(vessel_name).upper(),
        'X1_IMO': vessel_imo,
        'X1_VSLSD': supply_dates,
        'X1_MQ': mgo_tons,
        'X1_MP': mgo_price - 2,
        'X1_AGNT': agent,
    }
    replacements2 = {
        'X1_RN': str(get_bunker_date(replacements.get('X1_VSLSD')).strftime('%Y%m%d')) + '-NOM-' + str(replacements.get('X1_VSLN')).replace(' ', '_'),
        'X1_UC': 'USD',
        'X1_DDD': 30,
        'X1_DATE': get_bunker_date(replacements.get('X1_VSLSD')) - timedelta(days=10),
    }
    in_path = MGO_TEMPLATE
    out_path = os.path.join(FINISHED_DIR, f"{replacements2['X1_RN']}.docx")
    replace_strings_in_docx(in_path, out_path, replacements, 1)
    replace_strings_in_docx(out_path, out_path, replacements2, 1)
    input_docx = out_path
    output_pdf = os.path.join(FINISHED_DIR, f"{replacements2['X1_RN']}.pdf")
    final_path = convert_docx_to_pdf(input_docx, output_pdf, LIBREOFFICE_PATH)
    if final_path.endswith('.pdf'):
        delete_file(out_path)
    queued_up_files.append(final_path)


def process_ifo(vessel_name, vessel_imo, supply_dates, ifo_tons, ifo_price, agent):
    global queued_up_files
    if not os.path.exists(IFO_TEMPLATE):
        tmp = Document()
        tmp.add_heading('Bunkering nomination (IFO)', level=1)
        tmp.add_paragraph(f'Vessel: {vessel_name} (IMO {vessel_imo})')
        tmp.add_paragraph(f'Date: {supply_dates}')
        tmp.save(IFO_TEMPLATE)
    replacements = {
        'X1_CN': str('Simple Fuel FZCO').upper(),
        'X1_VSLN': str(vessel_name).upper(),
        'X1_IMO': vessel_imo,
        'X1_VSLSD': supply_dates,
        'X1_IQ': ifo_tons,
        'X1_IP': ifo_price - 2,
        'X1_AGNT': agent,
    }
    replacements2 = {
        'X1_RN': str(get_bunker_date(replacements.get('X1_VSLSD')).strftime('%Y%m%d')) + '-NOM-' + str(replacements.get('X1_VSLN')).replace(' ', '_'),
        'X1_UC': 'USD',
        'X1_DDD': 30,
        'X1_DATE': get_bunker_date(replacements.get('X1_VSLSD')) - timedelta(days=10),
    }
    in_path = IFO_TEMPLATE
    out_path = os.path.join(FINISHED_DIR, f"{replacements2['X1_RN']}.docx")
    replace_strings_in_docx(in_path, out_path, replacements, 1)
    replace_strings_in_docx(out_path, out_path, replacements2, 1)
    input_docx = out_path
    output_pdf = os.path.join(FINISHED_DIR, f"{replacements2['X1_RN']}.pdf")
    final_path = convert_docx_to_pdf(input_docx, output_pdf, LIBREOFFICE_PATH)
    if final_path.endswith('.pdf'):
        delete_file(out_path)
    queued_up_files.append(final_path)


app = FastAPI()

# CORS for local dev (Next.js at http://localhost:3000 etc.)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

class get_nom_info(BaseModel):
    vessel_name: str | None = ""
    vessel_imo: int | None = 0
    vessel_port: str | None = ""
    mgo_tons: str | None = "0"
    mgo_price: float | None = 0.0
    ifo_tons: str | None = "0"
    ifo_price: float | None = 0.0
    vessel_supply_date: str | None = ""
    vessel_trader: str | None = ""
    vessel_agent: str | None = ""


@app.get('/')
def root():
    return {'msg': 'Welcome to the API'}


@app.get('/download/{filename}')
async def download_file(filename: str):
    """Download a generated file"""
    file_path = os.path.join(FINISHED_DIR, filename)
    if not os.path.exists(file_path):
        return {'error': 'File not found'}, 404
    return FileResponse(
        file_path,
        media_type='application/pdf' if filename.endswith('.pdf') else 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename=filename
    )


def process_noms(full_vessel_data):
    global queued_up_files
    queued_up_files = []

    if (full_vessel_data['mgo_tons'] != '0') and (full_vessel_data['ifo_tons'] == '0'):
        process_mgo(
            vessel_name=full_vessel_data['vessel_name'],
            vessel_imo=full_vessel_data['vessel_imo'],
            supply_dates=full_vessel_data['vessel_supply_date'],
            mgo_tons=full_vessel_data['mgo_tons'],
            mgo_price=full_vessel_data['mgo_price'],
            agent=full_vessel_data['vessel_agent'],
        )
    elif (full_vessel_data['mgo_tons'] == '0') and (full_vessel_data['ifo_tons'] != '0'):
        process_ifo(
            vessel_name=full_vessel_data['vessel_name'],
            vessel_imo=full_vessel_data['vessel_imo'],
            supply_dates=full_vessel_data['vessel_supply_date'],
            ifo_tons=full_vessel_data['ifo_tons'],
            ifo_price=full_vessel_data['ifo_price'],
            agent=full_vessel_data['vessel_agent'],
        )
    elif (full_vessel_data['mgo_tons'] != '0') and (full_vessel_data['ifo_tons'] != '0'):
        process_both(
            vessel_name=full_vessel_data['vessel_name'],
            vessel_imo=full_vessel_data['vessel_imo'],
            supply_dates=full_vessel_data['vessel_supply_date'],
            mgo_tons=full_vessel_data['mgo_tons'],
            mgo_price=full_vessel_data['mgo_price'],
            ifo_tons=full_vessel_data['ifo_tons'],
            ifo_price=full_vessel_data['ifo_price'],
            agent=full_vessel_data['vessel_agent'],
        )

    fetched_email_body = (
        f"Dear Simple Fuel FZCO,\n\nKindly find the below attached nomination for the vessel:\n\n"
        f"- mv {full_vessel_data['vessel_name']} : IMO {full_vessel_data['vessel_imo']} : supply on {full_vessel_data['vessel_supply_date']}\n\n"
        f"Warm regards"
    )
    fetched_email_subject = f"NOMINATION FOR VESSEL: {full_vessel_data['vessel_name']} (IMO: {full_vessel_data['vessel_imo']})"
    send_email(recipients=[PEN_EMAIL], subject=fetched_email_subject, body=fetched_email_body, attachments=queued_up_files)

    s3_files = upload_files_to_s3(queued_up_files)
    return {
        'local_files': queued_up_files,
        's3_files': s3_files,
    }


@app.post('/endpoint1')
async def endpoint1(item: get_nom_info):
    vessel_name = str(item.vessel_name)
    vessel_imo = int(item.vessel_imo)
    vessel_port = str(item.vessel_port)
    mgo_tons = str(item.mgo_tons)
    mgo_price = float(item.mgo_price)
    ifo_tons = str(item.ifo_tons)
    ifo_price = float(item.ifo_price)
    vessel_supply_date = str(item.vessel_supply_date)
    vessel_trader = str(item.vessel_trader)
    vessel_agent = str(item.vessel_agent)

    nomination_data = {
        'vessel_name': vessel_name,
        'vessel_imo': vessel_imo,
        'vessel_port': vessel_port,
        'mgo_tons': mgo_tons,
        'mgo_price': mgo_price,
        'ifo_tons': ifo_tons,
        'ifo_price': ifo_price,
        'vessel_supply_date': vessel_supply_date,
        'vessel_trader': vessel_trader,
        'vessel_agent': vessel_agent
    }

    print(nomination_data)
    result = process_noms(nomination_data)
    return {'ok': True, 'received': nomination_data, 'files': result.get('s3_files') or [], 'local_files': result.get('local_files') or []}


class InvoiceData(BaseModel):
    vessel_name: str
    vessel_imo: int
    vessel_flag: str
    vessel_port: str
    bdn_numbers: str
    mgo_tons: str = "0"
    mgo_price: float = 0.0
    ifo_tons: str = "0"
    ifo_price: float = 0.0
    supply_date: str
    currency: str = "USD"
    exchange_rate: float = 1.0
    company_name: str = "Simple Fuel FZCO"
    company_address: str = ""


def determine_bank(currency):
    """Determine bank details based on currency"""
    bank_details = {
        'AED': {
            'name': 'Bank of Dubai',
            'account': '10000000000001',
            'iban': 'AE10000000000001',
            'swift': 'AEDABUXXX'
        },
        'USD': {
            'name': 'Bank of Dubai',
            'account': '10000000000002',
            'iban': 'AE10000000000002',
            'swift': 'AEDABUXXX'
        },
        'EUR': {
            'name': 'Bank of Dubai',
            'account': '10000000000003',
            'iban': 'AE10000000000003',
            'swift': 'AEDABUXXX'
        },
        'BHD': {
            'name': 'Bank of Dubai',
            'account': '10000000000004',
            'iban': 'AE10000000000004',
            'swift': 'AEDABUXXX'
        }
    }
    return bank_details.get(currency, bank_details['USD'])


@app.post('/generate-invoice')
async def generate_invoice(invoice_data: InvoiceData):
    """Generate invoice PDF and return file path or S3 URL"""
    global queued_up_files
    queued_up_files = []
    
    try:
        # Determine which template to use based on product types
        has_mgo = invoice_data.mgo_tons and invoice_data.mgo_tons != "0"
        has_ifo = invoice_data.ifo_tons and invoice_data.ifo_tons != "0"
        
        # Use invoice templates if they exist, otherwise fallback to nomination templates
        if has_mgo and has_ifo:
            template_path = BOTH_INVOICE_TEMPLATE if os.path.exists(BOTH_INVOICE_TEMPLATE) else BOTH_TEMPLATE
        elif has_mgo:
            template_path = MGO_INVOICE_TEMPLATE if os.path.exists(MGO_INVOICE_TEMPLATE) else MGO_TEMPLATE 
        elif has_ifo:
            template_path = IFO_INVOICE_TEMPLATE if os.path.exists(IFO_INVOICE_TEMPLATE) else IFO_TEMPLATE
        else:
            return {'ok': False, 'error': 'No products selected'}
        
        # Calculate totals
        mgo_total = float(invoice_data.mgo_tons or 0) * invoice_data.mgo_price * invoice_data.exchange_rate if has_mgo else 0
        ifo_total = float(invoice_data.ifo_tons or 0) * invoice_data.ifo_price * invoice_data.exchange_rate if has_ifo else 0
        subtotal = mgo_total + ifo_total
        
        # Get bank details
        bank = determine_bank(invoice_data.currency)
        
        # Prepare replacements
        replacements = {
            'X1_CN': invoice_data.company_name.upper(),
            'X1_COADR': invoice_data.company_address.upper(),
            'X1_UC': invoice_data.currency,
            'X1_UXER': invoice_data.exchange_rate,
            'X1_VSLN': invoice_data.vessel_name.upper(),
            'X1_IMO': invoice_data.vessel_imo,
            'X1_VSLF': invoice_data.vessel_flag.upper(),
            'X1_VSLP': invoice_data.vessel_port.upper(),
            'X1_VSLSD': invoice_data.supply_date,
            'X1_BDN': invoice_data.bdn_numbers,
        }
        
        if has_mgo:
            replacements['X1_MQ'] = invoice_data.mgo_tons
            replacements['X1_MP'] = invoice_data.mgo_price
            replacements['X1_MGOT'] = f"{mgo_total:,.2f}"
            
        if has_ifo:
            replacements['X1_IQ'] = invoice_data.ifo_tons
            replacements['X1_IP'] = invoice_data.ifo_price
            replacements['X1_IFOT'] = f"{ifo_total:,.2f}"
        
        # Additional replacements
        supply_date_obj = get_bunker_date(invoice_data.supply_date)
        payment_deadline = supply_date_obj + timedelta(days=10)
        
        replacements2 = {
            'X1_TOTAL': f"{subtotal:,.2f}",
            'X1_SBTTL': f"{subtotal:,.2f}",
            'X1_PYMTD': payment_deadline,
            'X1_DATE': supply_date_obj,
            'X1_RN': supply_date_obj.strftime("%Y%m%d") + '-INV-' + invoice_data.vessel_name.replace(' ', '_'),
            'X2_BANK': bank['name'].upper(),
            'X2_SWIFT': bank['swift'],
            'X2_ULIBAN': bank['iban'],
            'X2_ULAN': bank['account']
        }
        
        # Generate files
        if not os.path.exists(template_path):
            return {'ok': False, 'error': f'Template not found: {template_path}'}
            
        output_filename = replacements2['X1_RN']
        out_path = os.path.join(FINISHED_DIR, f"{output_filename}.docx")
        
        replace_strings_in_docx(template_path, out_path, replacements, 1)
        replace_strings_in_docx(out_path, out_path, replacements2, 1)
        
        # Convert to PDF
        output_pdf = os.path.join(FINISHED_DIR, f"{output_filename}.pdf")
        final_path = convert_docx_to_pdf(out_path, output_pdf, LIBREOFFICE_PATH)
        
        # Delete intermediate DOCX if PDF was created
        if final_path.endswith('.pdf'):
            delete_file(out_path)
            
        queued_up_files.append(final_path)
        
        # Upload to S3 if configured
        s3_files = upload_files_to_s3(queued_up_files)
        
        return {
            'ok': True,
            'message': 'Invoice generated successfully',
            'local_files': queued_up_files,
            's3_files': s3_files,
            'filename': os.path.basename(final_path)
        }
    except Exception as e:
        return {'ok': False, 'error': str(e), 'message': f'Failed to generate invoice: {str(e)}'}


class InitialRequest(BaseModel):
    vessel_name: str
    mgo_tons: str
    ifo_tons: str
    bunker_date_start: str
    bunker_date_end: str
    port: str
    agent_name: str
    full_order_text: str


@app.post('/initial-request')
async def initial_request(request_data: InitialRequest):
    """Send initial bunker request email"""
    try:
        # Compose email with request details
        email_body = f"""Dear Simple Fuel FZCO,

We would like to request a bunker nomination for:

{request_data.full_order_text}

Best regards"""
        
        email_subject = f"BUNKER REQUEST - {request_data.vessel_name} at {request_data.port}"
        
        # Send email (will skip if DISABLE_EMAIL=1 or token missing)
        send_email(
            recipients=[PEN_EMAIL],
            subject=email_subject,
            body=email_body,
            attachments=[]
        )
        
        return {
            'ok': True,
            'message': 'Initial request email sent successfully',
            'data': request_data.dict()
        }
    except Exception as e:
        return {'ok': False, 'error': str(e), 'message': f'Failed to send email: {str(e)}'}


class FirstNominationData(BaseModel):
    vessel_name: str
    vessel_imo: int
    vessel_flag: str


@app.post('/first-nomination')
async def first_nomination(nom_data: FirstNominationData):
    """Generate and email first nomination with basic vessel info"""
    try:
        email_body = f"""Dear Simple Fuel FZCO,

Please find our first nomination for:

Vessel: {nom_data.vessel_name}
IMO: {nom_data.vessel_imo}
Flag: {nom_data.vessel_flag}

We will provide detailed quantities and dates in the final nomination.

Best regards"""
        
        email_subject = f"FIRST NOMINATION - {nom_data.vessel_name} (IMO: {nom_data.vessel_imo})"
        
        send_email(
            recipients=[PEN_EMAIL],
            subject=email_subject,
            body=email_body,
            attachments=[]
        )
        
        return {
            'ok': True,
            'message': 'First nomination sent successfully',
            'data': nom_data.dict()
        }
    except Exception as e:
        return {'ok': False, 'error': str(e), 'message': f'Failed to send: {str(e)}'}


class FinalNominationRequest(BaseModel):
    vessel_name: str
    actual_mgo_tons: str
    mgo_price: float
    actual_ifo_tons: str
    ifo_price: float
    bunker_date: str


@app.post('/final-nomination')
async def final_nomination(nom_data: FinalNominationRequest):
    """Generate final nomination with actual quantities and send email with PDF"""
    try:
        email_body = f"""Dear Simple Fuel FZCO,

Please find our final nomination:

Vessel: {nom_data.vessel_name}
Supply Date: {nom_data.bunker_date}

Products:
- MGO: {nom_data.actual_mgo_tons} MT @ USD {nom_data.mgo_price}/MT
- IFO: {nom_data.actual_ifo_tons} MT @ USD {nom_data.ifo_price}/MT

Best regards"""
        
        email_subject = f"FINAL NOMINATION - {nom_data.vessel_name}"
        
        send_email(
            recipients=[PEN_EMAIL],
            subject=email_subject,
            body=email_body,
            attachments=[]
        )
        
        return {
            'ok': True,
            'message': 'Final nomination sent successfully',
            'data': nom_data.dict()
        }
    except Exception as e:
        return {'ok': False, 'error': str(e), 'message': f'Failed to send: {str(e)}'}


if __name__ == '__main__':
    import uvicorn
    uvicorn.run(app, host='0.0.0.0', port=8000)


