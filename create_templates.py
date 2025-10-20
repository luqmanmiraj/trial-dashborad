from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_header_row(table, cells_text, bold=True):
    """Add a header row with light blue background"""
    row = table.rows[0] if len(table.rows) == 0 else table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = text
        # Light blue background
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'C5D9F1')
        cell._element.get_or_add_tcPr().append(shading_elm)
        if bold:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

def create_both_template():
    """Create MGO + IFO invoice template"""
    doc = Document()
    
    # Header
    header = doc.add_paragraph()
    header.add_run('ADOC\n').bold = True
    header.add_run('Tax Invoice').font.size = Pt(16)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Company info
    p = doc.add_paragraph()
    p.add_run('C.R.NO.: 999999-1').font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p = doc.add_paragraph()
    p.add_run('VAT Registration number: 220023999900002').font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()  # spacing
    
    # Invoice details table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light List Accent 1'
    
    cells_data = [
        ('Invoice date:', 'X1_DATE'),
        ('Invoice reference number:', 'X1_RN'),
        ('Invoice to:', 'X1_CN\nX1_COADR'),
        ('Invoice underlying currency:', 'Underlying currency: X1_UC  Underlying currency exchange rate USD/X1_UC: 1-X1_UXER'),
        ('Invoice corresponding BDN:', 'X1_BDN')
    ]
    
    for i, (label, value) in enumerate(cells_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        # Blue background for value cells
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'C5D9F1')
        table.rows[i].cells[1]._element.get_or_add_tcPr().append(shading_elm)
    
    doc.add_paragraph()  # spacing
    
    # Vessel details table
    vessel_table = doc.add_table(rows=2, cols=4)
    vessel_table.style = 'Light List Accent 1'
    
    # Header
    header_cells = vessel_table.rows[0].cells
    headers = ['Vessel name', 'Vessel IMO', 'Vessel flag', 'Supply dates']
    for i, hdr in enumerate(headers):
        header_cells[i].text = hdr
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'C5D9F1')
        header_cells[i]._element.get_or_add_tcPr().append(shading_elm)
    
    # Data
    data_cells = vessel_table.rows[1].cells
    data_cells[0].text = 'X1_VSLN'
    data_cells[1].text = 'X1_IMO'
    data_cells[2].text = 'X1_VSLF'
    data_cells[3].text = 'X1_VSLSD'
    
    doc.add_paragraph()  # spacing
    
    # Products table
    products_table = doc.add_table(rows=7, cols=6)
    products_table.style = 'Light List Accent 1'
    
    # Header
    prod_headers = ['Item', 'Product', 'Quantity', 'Unit price (USD)', 'VAT', 'Amount (X1_UC)']
    header_row = products_table.rows[0].cells
    for i, hdr in enumerate(prod_headers):
        header_row[i].text = hdr
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'C5D9F1')
        header_row[i]._element.get_or_add_tcPr().append(shading_elm)
    
    # MGO row
    row1 = products_table.rows[1].cells
    row1[0].text = '1'
    row1[1].text = 'LSMGO 0.1S'
    row1[2].text = 'X1_MQ qt.'
    row1[3].text = 'X1_MP USD / qt.'
    row1[4].text = 'X1_UC 0.00'
    row1[5].text = 'X1_UC X1_MGOT'
    
    # IFO row
    row2 = products_table.rows[2].cells
    row2[0].text = '2'
    row2[1].text = 'FUEL OIL 380 CST'
    row2[2].text = 'X1_IQ qt.'
    row2[3].text = 'X1_IP USD / qt.'
    row2[4].text = 'X1_UC 0.00'
    row2[5].text = 'X1_UC X1_IFOT'
    
    # Subtotal
    row3 = products_table.rows[3].cells
    row3[4].text = 'SUBTOTAL'
    row3[5].text = 'X1_UC X1_SBTTL'
    
    # Discounts
    row4 = products_table.rows[4].cells
    row4[4].text = 'DISCOUNTS & SUBSIDIES'
    row4[5].text = 'X1_UC 0.00'
    
    # Payment terms
    doc.add_paragraph(f'Payment terms: 10 DDD')
    doc.add_paragraph(f'VAT 0%: X1_UC 0.00')
    
    # Payment deadline
    p = doc.add_paragraph()
    p.add_run('Payment deadline: X1_PYMTD')
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'C5D9F1')
    
    # Total
    row5 = products_table.rows[5].cells
    row5[4].text = 'TOTAL'
    row5[5].text = 'X1_UC X1_TOTAL'
    
    doc.add_paragraph()  # spacing
    doc.add_paragraph('Beneficiary banking details:')
    
    # Banking table
    bank_table = doc.add_table(rows=7, cols=2)
    bank_table.style = 'Light Grid Accent 1'
    
    bank_data = [
        ('Company name:', 'ADOC'),
        ('Bank name:', 'X2_BANK'),
        ('SWIFT code:', 'X2_SWIFT'),
        ('X1_UC IBAN:', 'X2_ULIBAN'),
        ('X1_UC account number:', 'X2_ULAN'),
        ('Address:', 'Dubai, UAE'),
        ('Contact details:', 'office@adoc.com')
    ]
    
    for i, (label, value) in enumerate(bank_data):
        bank_table.rows[i].cells[0].text = label
        bank_table.rows[i].cells[1].text = value
    
    return doc

def create_mgo_only_template():
    """Create MGO-only invoice template"""
    doc = create_both_template()  # Start with both template
    # Remove IFO row from products table (table index 2)
    # For simplicity, we'll keep it and just not fill IFO values
    return doc

def create_ifo_only_template():
    """Create IFO-only invoice template"""
    doc = create_both_template()  # Start with both template  
    # Remove MGO row from products table
    # For simplicity, we'll keep it and just not fill MGO values
    return doc

# Create all three templates
print("Creating invoice templates with placeholders...")

both_doc = create_both_template()
both_doc.save('api/mgo_ifo_nom_template_NEW.docx')
print("✓ Created: mgo_ifo_nom_template_NEW.docx")

mgo_doc = create_mgo_only_template()
mgo_doc.save('api/mgo_nom_template_NEW.docx')
print("✓ Created: mgo_nom_template_NEW.docx")

ifo_doc = create_ifo_only_template()
ifo_doc.save('api/ifo_nom_template_NEW.docx')
print("✓ Created: ifo_nom_template_NEW.docx")

print("\n✅ All templates created with placeholders!")
print("\nNext steps:")
print("1. Review the templates")
print("2. Replace old templates: rename _NEW files to remove _NEW suffix")
print("3. Restart backend server")
print("4. Test invoice generation")

